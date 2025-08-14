import os
import docx
import re
import pdfplumber
import google.generativeai as genai
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet
from pdf2docx import Converter
import streamlit as st
import pandas as pd  
from io import BytesIO
import tempfile
import zipfile
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.styles import ParagraphStyle


pdfmetrics.registerFont(TTFont('Algerian', 'ALGER.TTF'))
pdfmetrics.registerFont(TTFont('TimesNewRoman', 'C:/Windows/Fonts/times.ttf'))

# Google Gemini API configuration
genai.configure(api_key="AIzaSyDx1y2D6_n1onaUrjJQxCJ58LeE8wwDZRo")
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 39,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain"
}

# Function to convert PDF to DOCX
def convert_pdf_to_docx(pdf_file_bytes, docx_buffer):
    """Converts a PDF file to a Word document (DOCX) format."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
        temp_pdf.write(pdf_file_bytes)
        temp_pdf_path = temp_pdf.name

    cv = Converter(temp_pdf_path)
    cv.convert(docx_buffer)
    cv.close()

    os.remove(temp_pdf_path)

# Function to extract text from Word document (docx)
def extract_text_from_word(file):
    doc = docx.Document(file)
    full_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(row_data))

    return "\n".join(full_text)

# Function to extract text from PDF document
def extract_text_from_pdf(file):
    full_text = []

    with pdfplumber.open(file) as pdf:
        # Extract text from each page
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                full_text.append(page_text.strip())

            # Extract tables
            for table in page.extract_tables():
                for row in table:
                    row_data = [str(cell).strip() if cell else '' for cell in row]
                    full_text.append(" | ".join(row_data))

    return "\n".join(full_text)

# Function to automatically detect file type and extract text
def extract_text(file):
    file_extension = os.path.splitext(file.name)[1].lower()
    if file_extension == '.docx':
        return extract_text_from_word(file)
    elif file_extension == '.pdf':
        return extract_text_from_pdf(file)
    else:
        raise ValueError("Unsupported file type. Please provide a .docx or .pdf file.")

# Function to create the PDF question paper
def create_question_paper( subject, subject_code, comp_question, department,
                          date, questions, pdf_file,exam_name ):
    document = SimpleDocTemplate(pdf_file, pagesize=letter, topMargin=20, bottomMargin=30)
    styles = getSampleStyleSheet()
    story = []
    #new strt
    small_style = styles['Normal'].clone('SmallStyle')
    small_style.fontSize = 14
    small_style.alignment = 1  # 0 = left, 1 = center, 2 = right
    small_style.fontName = 'Times-BoldItalic'
    
    # Add this where you want to include the image
    story.append(Spacer(1, 8))
    
    image_path = r"ifetname.jpg"  # Change to your image path
    image = Image(image_path)
    # Set the desired width (maintain aspect ratio)
    desired_width = 600  # Change to your preferred width
    image.drawHeight = desired_width * (image.drawHeight / image.drawWidth)
    image.drawWidth = desired_width
    # Create a table to position the image in the right corner
    header_data = [
        [Spacer(1, 0), image]  # Spacer to push the image to the right
    ]
    header_table = Table(header_data, colWidths=[None, None])  # Adjust as needed
    header_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
    ]))

    story.append(header_table)
    
    #new end

    
    
    styles['Heading2'].alignment = TA_CENTER
    story.append(Paragraph('<font name="Algerian" size="16" color="black"><i>{}</i></font>'.format(f"{exam_name}"), styles['Heading2']))
    story.append(Spacer(1, 8))

    table_text_style = ParagraphStyle(
        name='TableTextStyle',
        fontName='Times-Roman',
        fontSize=11,
        textColor=colors.black,
        leading=14  # Line spacing
    )

    # Header data with left and right alignment for each row
    header_data = [
        [
            Paragraph(f"SUBJECT CODE: {subject_code}", table_text_style),  # Left-aligned
            Paragraph(f"MAX MARKS: 100", table_text_style)  # Right-aligned
        ],
        [
            Paragraph(f"SUBJECT NAME: {subject}", table_text_style),  # Left-aligned
            Paragraph(f"DURATION: 180Mins", table_text_style)  # Right-aligned
        ],
        [
            Paragraph(f"DATE: {date}", table_text_style),  # Left-aligned
            Paragraph(f"YEAR/SEMESTER: III/V", table_text_style)  # Right-aligned
        ]
    ]

    # Create the table
    header_table = Table(header_data, colWidths=['*', '*'])  # Equal column widths
    header_table.setStyle(TableStyle([
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),  # Left align the first column
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),  # Right align the second column
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Vertical alignment at the top
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ('TOPPADDING', (0, 0), (-1, -1), 3),  # Top padding
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),  # Bottom padding
    ]))

    # Add the table and a spacer to the story
    story.append(header_table)
    story.append(Spacer(1, 6))

    # Add Part A Questions
    story.append(Spacer(1, 16))
    story.append(Paragraph('<font name="Times-Bold" size="12">{}</font>'.format("Part A"), styles['Title']))
    story.append(Spacer(1, 2))

    custom_style = ParagraphStyle(
    name='CustomStyle',
    fontName='Times-Roman',  # Set to Times New Roman
    fontSize=12,  # Set font size to 12
    textColor=colors.black,  # Set text color to black
    leading=14  # Optional: line height
)
  
    # Instructions for the exam
    story.append(Paragraph("Answer ALL Questions", small_style))
    story.append(Spacer(1, 6))
    for question in questions['Part A']:
        question = question.strip().replace('#', '')
        if question:
            story.append(Paragraph(question, custom_style))
            story.append(Spacer(1, 8))
  # Add Part B Questions
    story.append(Spacer(1, 2))
    story.append(Paragraph('<font name="Times-Bold" size="12">{}</font>'.format("Part B"), styles['Title']))
    story.append(Spacer(1, 2))
    for question in questions['Part B']:
        question = question.strip().replace('#', '')
        if question:
            story.append(Paragraph(question, custom_style))
            story.append(Spacer(1, 10))

    # Build the PDF
    document.build(story)

# Streamlit application
def main():
    st.title("Question Paper Generator")

    # File upload for DOCX or PDF
    uploaded_file = st.file_uploader("Upload a Word or PDF document  upto 5 files", type=["docx", "pdf"], accept_multiple_files=True,)
    if uploaded_file and len(uploaded_file) > 5:
        st.error("You can upload a maximum of 5 files.")
        return
    extracted_text = ""
    if uploaded_file:
        for uploaded_file in uploaded_file:
            try:
                file_text = extract_text(uploaded_file)
                extracted_text += f"\n{file_text}\n"  # Combine text from all files
            except ValueError as e:
                st.error(f"Error processing {uploaded_file.name}: {e}")

        # Display extracted text in a text area
        st.text_area("Extracted Text", extracted_text, height=200)

        # Collect other information
        exam_name = st.text_input("exam name", "IA Examination")
        subject = st.text_input("Subject", "Introduction To Machine Learning")
        subject_code = st.text_input("Subject Code", "19UADPC501")
        comp_question = st.text_input("Enter the questions that must be included as compulsory in the paper: E.g., Q1: What is AI?  ", "")
        department = st.text_input("Department", "Artificial Intellinence And Data Science")
        date = st.date_input("Date", value=pd.Timestamp.now().date())

        # Difficulty level selection
        difficulty_level = st.radio("Select Difficulty Level", ["Easy", "Medium", "Hard"])


        # Filename inputs for both PDF and DOCX files
        pdf_filename = st.text_input("PDF Filename", "question_paper")
        docx_filename = st.text_input("DOCX Filename", "question_paper")

        # Ensure the extensions are added correctly
        if not pdf_filename.endswith(".pdf"):
            pdf_filename += ".pdf"

        if not docx_filename.endswith(".docx"):
            docx_filename += ".docx"

        # Generate the question paper based on API response
        if st.button("Generate and Download PDF & DOCX"):
            # Generate questions only once
            questions = generate_questions(extracted_text, difficulty_level,comp_question)

            # Create PDF in-memory
            pdf_file = BytesIO()
            
            create_question_paper(
            subject=subject,
            subject_code=subject_code,
            comp_question=comp_question,
            department=department,
            date=date,
            questions=questions,
            pdf_file=pdf_file,
            exam_name=exam_name

        )

            pdf_file.seek(0)

            # Convert the PDF to DOCX
            docx_file = BytesIO()
            convert_pdf_to_docx(pdf_file.getvalue(), docx_file)

            # Zip the PDF and DOCX files together for single download
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zf:
                zf.writestr(pdf_filename, pdf_file.getvalue())
                zf.writestr(docx_filename, docx_file.getvalue())

            zip_buffer.seek(0)
            st.download_button("Download PDF & DOCX", data=zip_buffer.getvalue(), file_name="question_paper.zip", mime="application/zip")


# Modify the generate_questions function to include difficulty level
def generate_questions(extracted_text, difficulty_level, comp_question):
    model = genai.GenerativeModel(model_name="gemini-1.5-flash", generation_config=generation_config)
    chat_session = model.start_chat(history=[])

    # Send the message to the model to process the extracted text with difficulty level
    response = chat_session.send_message(
    f"""
    I have extracted text from a document that contains questions divided into two distinct parts: "Part A" and "Part B." The questions are structured as follows:
    
    1. In "Part A," questions are direct and carry equal marks.
    2. In "Part B," questions can either:
       - Be a single complete question worth 16 marks.
       - Be split into two sub-parts labeled "(i)" and "(ii)," where each sub-part is worth 8 marks. However, these sub-parts should be treated as one complete question collectively worth 16 marks.

    Task Instructions:
    1. For "Part A":
       - Select exactly 10 questions from the extracted text based on the specified difficulty level '{difficulty_level}'.
       - Ensure the selection is random but aligned with the provided difficulty level.
       - Each question should appear on a new line, numbered sequentially, strictly adhering to this format:
       
       Part A:
       1. Question 1
       2. Question 2
       ...
       10. Question 10

    2. For "Part B":
       - Generate exactly 5 questions, each comprising two options (a and b), from the extracted text based on the specified difficulty level '{difficulty_level}'.
       - Format each question pair as follows:
         - Label the first option with "1.a)" and write the question.
         - On the next line, write "(or)" without any additional text, symbols, or characters.
         - Label the second option with "1.b)" and write the question.
       - Repeat this format for all five question pairs as shown below:

       Part B:
       1.a Question 1a
       (or)
       1.b Question 1b
       2.a Question 2a
       (or)
       2.b Question 2b
       ...
       5.a Question 5a
       (or)
       5.b Question 5b

    Formatting Guidelines:
    - Do not include any additional text, symbols, or explanations such as ":, *, **, random selection," or similar.
    - Maintain clarity, simplicity, and strict adherence to the format provided.
    - Avoid redundancy or unnecessary commentary in the output.
    - Do not give question number like 1.a) just give like 1.a

    Special Requirements:
    - Compulsory Question: If the specified compulsory question '{comp_question}' is present in the extracted text, ensure it is included in the corresponding part (Part A or Part B) based on its original categorization. If it is absent from the extracted text, do not generate it.
    - When including the compulsory question, ensure it is integrated seamlessly into the format without any extra labels or explanations.

    Extracted Text:
    {extracted_text}
    """
)


    # # Split response into Part A and Part B
    # parts = response.text.split("Part B")
    # part_a_questions = parts[0].replace("Part A", "").strip().split("\n")
    # part_b_questions = parts[1].strip().split("\n") if len(parts) > 1 else []

    # return {
    #     'Part A': [q for q in part_a_questions if q and "Questions:" not in q],
    #     'Part B': [q for q in part_b_questions if q and "Questions:" not in q]
    # }
    
    def clean_question_text(question):
        """
        Cleans up a single question by removing unwanted symbols, phrases, and extra spaces.
        """
        # Define unwanted patterns
        unwanted_patterns = [
            r"\*{1,2}",       # Remove * or ** symbols
            r"\s*[:;]\s*",    # Remove colons and semicolons with surrounding spaces
            r"randomly selected questions",  # Remove the phrase 'randomly selected questions'
            #r"^\d+\.",         # Remove numbering like '1.', '2.', etc. (if needed)
            #r"^\d+\.\w\)",               # Change '1.a)' to '1.a'
            r"\s{2,}",         # Remove extra spaces
            r"^\s+|\s+?$"      # Trim leading/trailing spaces
        ]
        for pattern in unwanted_patterns:
            question = re.sub(pattern, "", question)
        return question.strip()

    # Split response into Part A and Part B
    parts = response.text.split("Part B")
    part_a_questions = parts[0].replace("Part A", "").strip().split("\n")
    part_b_questions = parts[1].strip().split("\n") if len(parts) > 1 else []

    # Clean and filter questions for both parts
    return {
        'Part A': [clean_question_text(q) for q in part_a_questions if q and "Questions:" not in q],
        'Part B': [clean_question_text(q) for q in part_b_questions if q and "Questions:" not in q]
    }


if __name__ == "__main__":
    main()
